from __future__ import annotations

import shutil
import tempfile
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "Khung-tai-lieu-bao-cao-do-an.docx"
REPORT = ROOT / "docs" / "BAO_CAO_DO_AN_CHUYEN_NGANH_CODELEARN.docx"
ASSET_DIR = ROOT / "docs" / "report-assets"
TEMP_REPORT = REPORT.with_name(REPORT.stem + ".tmp.docx")

GREEN = "009B72"
BLACK = "000000"
WHITE = "FFFFFF"
LIGHT_GREEN = "E7F7F1"
LIGHT_GRAY = "F2F4F7"
RED = "C2415D"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Nhấn F9 để cập nhật trường tự động."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, placeholder, end])


def set_run_font(run, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.2)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.3
    normal.paragraph_format.space_after = Pt(6)

    for name, size in [
        ("Heading 1", 16),
        ("Heading 2", 14),
        ("Heading 3", 13),
    ]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLACK)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    if "Code Block" not in [style.name for style in doc.styles]:
        code = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = doc.styles["Code Block"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code.font.size = Pt(9)
    code.paragraph_format.left_indent = Cm(0.6)
    code.paragraph_format.right_indent = Cm(0.4)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(6)

    caption = doc.styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    caption.font.size = Pt(12)
    caption.font.bold = True
    caption.font.italic = False
    caption.font.color.rgb = RGBColor.from_string(BLACK)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.keep_with_next = True

    for sec in doc.sections:
        header = sec.header.paragraphs[0]
        header.text = "CODELEARN - TÀI LIỆU THIẾT KẾ CHI TIẾT HỆ THỐNG"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header.runs:
            set_run_font(run, 9)
            run.font.color.rgb = RGBColor.from_string(BLACK)
        footer = sec.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_field(footer, "PAGE")
        for run in footer.runs:
            set_run_font(run, 10)


def clear_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def add_paragraph(
    doc: Document,
    text: str = "",
    *,
    bold_prefix: str | None = None,
    indent: bool = True,
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    keep: bool = False,
):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = keep
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.0)
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, 13, True)
        rest = p.add_run(text[len(bold_prefix) :])
        set_run_font(rest, 13)
    else:
        run = p.add_run(text)
        set_run_font(run, 13)
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        set_run_font(run, 13)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        run = p.add_run(item)
        set_run_font(run, 13)


def add_code(doc: Document, code: str) -> None:
    p = doc.add_paragraph(style="Code Block")
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.keep_together = True
    set_cell_shading_proxy(p, "F7F8FA")
    run = p.add_run(code)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9)


def set_cell_shading_proxy(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


table_counter = 0
figure_counter = 0


def add_table(doc: Document, title: str, headers: list[str], rows: list[list[str]], widths=None):
    global table_counter
    table_counter += 1
    cap = doc.add_paragraph(style="Caption")
    cap.add_run(f"Bảng {table_counter}: {title}")
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    header = table.rows[0]
    set_repeat_table_header(header)
    set_row_cant_split(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        cell.text = value
        set_cell_shading(cell, WHITE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, 11, True)
                run.font.color.rgb = RGBColor.from_string(BLACK)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        for index, value in enumerate(values):
            cells[index].text = str(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_shading(cells[index], WHITE)
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.1
                for run in paragraph.runs:
                    set_run_font(run, 11)
        if widths:
            for index, width in enumerate(widths):
                cells[index].width = Cm(width)
    doc.add_paragraph()
    return table


def add_figure(doc: Document, image: Path, title: str, width_cm: float = 15.0) -> None:
    global figure_counter
    figure_counter += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(image), width=Cm(width_cm))
    cap = doc.add_paragraph(style="Caption")
    cap.add_run(f"Hình {figure_counter}: {title}")


def add_heading(doc: Document, text: str, level: int, page_break: bool = False) -> None:
    if page_break:
        doc.add_page_break()
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.page_break_before = False


def create_formula_figures(destination: Path) -> tuple[Path, Path]:
    plt.rcParams.update({"font.family": "DejaVu Sans", "font.size": 10})

    retention = destination / "retention-v3.png"
    labels = ["Gần nhất", "Độ đều 14 ngày", "Tiến độ", "Chất lượng", "Xu hướng tuần"]
    values = [25, 20, 20, 25, 10]
    colors = ["#009B72", "#20B486", "#4CB99A", "#146C94", "#C2415D"]
    fig, ax = plt.subplots(figsize=(10.5, 4.6), dpi=180)
    bars = ax.barh(labels[::-1], values[::-1], color=colors[::-1], height=0.62)
    for bar, value in zip(bars, values[::-1]):
        ax.text(value + 0.6, bar.get_y() + bar.get_height() / 2, f"{value} điểm", va="center", weight="bold")
    ax.set_xlim(0, 30)
    ax.set_xlabel("Điểm tối đa (tổng 100)")
    ax.set_title("RETENTION_V3_2026_07 - Điểm giữ nhịp ưu tiên hành vi gần đây", weight="bold", color="#172033")
    ax.grid(axis="x", alpha=0.2)
    for spine in ax.spines.values():
        spine.set_visible(False)
    fig.tight_layout()
    fig.savefig(retention, bbox_inches="tight", facecolor="white")
    plt.close(fig)

    competition = destination / "competition-v2.png"
    labels = ["Học bài", "Bài code", "Quiz", "Phòng thi", "Học đều"]
    values = [200, 250, 150, 300, 100]
    colors = ["#4CB99A", "#009B72", "#146C94", "#7A5AF8", "#C2415D"]
    fig, ax = plt.subplots(figsize=(10.5, 4.6), dpi=180)
    bars = ax.bar(labels, values, color=colors, width=0.62)
    for bar, value in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, value + 8, str(value), ha="center", weight="bold")
    ax.set_ylim(0, 350)
    ax.set_ylabel("Điểm tối đa")
    ax.set_title("COMPETITION_V2_1000 - Thang điểm thi đua chuẩn hóa", weight="bold", color="#172033")
    ax.grid(axis="y", alpha=0.2)
    for spine in ax.spines.values():
        spine.set_visible(False)
    fig.tight_layout()
    fig.savefig(competition, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return retention, competition


def add_cover(doc: Document) -> None:
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[TÊN TRƯỜNG]\n[TÊN KHOA/VIỆN]")
    set_run_font(run, 14, True)

    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("</> CodeLearn")
    set_run_font(run, 24, True)
    run.font.color.rgb = RGBColor.from_string(BLACK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TÀI LIỆU THIẾT KẾ CHI TIẾT HỆ THỐNG\n(Software Design Document - SDD)")
    set_run_font(run, 17, True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("XÂY DỰNG WEBSITE HỖ TRỢ HỌC LẬP TRÌNH CODELEARN")
    set_run_font(run, 18, True)
    run.font.color.rgb = RGBColor.from_string(BLACK)

    for _ in range(3):
        doc.add_paragraph()
    metadata = [
        "Nhóm/Sinh viên thực hiện: ........................................................................",
        "Mã sinh viên: .............................................................................................",
        "Giảng viên hướng dẫn: ................................................................................",
        "Lớp / Học kỳ: ...............................................................................................",
        "Ngày hoàn thành: 08/08/2026",
    ]
    for line in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(2.2)
        set_run_font(p.add_run(line), 13)

    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("[ĐỊA DANH], THÁNG 08 NĂM 2026"), 13, True)


def add_front_matter(doc: Document) -> None:
    add_heading(doc, "LỜI CAM ĐOAN", 1, page_break=True)
    add_paragraph(
        doc,
        "Nhóm em xin cam đoan tài liệu thiết kế chi tiết hệ thống CodeLearn là kết quả của quá trình khảo sát, phân tích, thiết kế, lập trình và kiểm thử do nhóm thực hiện. Các số liệu kiểm thử, cấu hình hạ tầng và giới hạn sản phẩm được trình bày theo bằng chứng hiện có trong mã nguồn; nhóm em không biến kết quả thử nghiệm trên máy demo thành cam kết cho môi trường production chưa được đo tải.",
    )
    add_paragraph(
        doc,
        "Nhóm em chịu trách nhiệm về tính chính xác của nội dung và sẵn sàng cung cấp repository, migration, dữ liệu seed, kết quả kiểm thử và kịch bản demo để giảng viên đối chiếu.",
    )

    add_heading(doc, "LỜI CẢM ƠN", 1, page_break=True)
    add_paragraph(
        doc,
        "Nhóm em xin chân thành cảm ơn giảng viên hướng dẫn đã góp ý để đề tài đi từ một website học lập trình thông thường đến một hệ thống có vòng kín giữ nhịp, chấm code tự chủ, AI Tutor local và dashboard can thiệp sớm. Các nhận xét về tính thực tế của công thức, khả năng mở rộng và cách chứng minh bằng dữ liệu đã giúp nhóm em hoàn thiện sản phẩm có trách nhiệm hơn.",
    )
    add_paragraph(
        doc,
        "Nhóm em cũng cảm ơn các thầy cô trong khoa đã trang bị kiến thức về phát triển web, cơ sở dữ liệu, phân tích thiết kế, an toàn thông tin và kiểm thử phần mềm. Đây là nền tảng để nhóm em hoàn thành tài liệu này.",
    )

    add_heading(doc, "MỤC LỤC", 1, page_break=True)
    p = doc.add_paragraph()
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u')

    add_heading(doc, "DANH SÁCH HÌNH VẼ, SƠ ĐỒ, BẢNG VÀ BIỂU ĐỒ", 1, page_break=True)
    p = doc.add_paragraph()
    add_field(p, 'TOC \\h \\z \\t "Caption,1"')

    add_heading(doc, "CÁC TASK ĐÃ THỰC HIỆN", 1, page_break=True)
    add_paragraph(
        doc,
        "Bảng dưới dùng để nhóm điền thông tin thành viên, MSSV và tỷ lệ đóng góp trước khi nộp chính thức. Phạm vi task đã được chia theo module để có thể đối chiếu với lịch sử Git.",
    )
    add_table(
        doc,
        "Phân công công việc của nhóm",
        ["Thành viên", "MSSV", "Nhiệm vụ chính", "Tỷ lệ", "Ghi chú"],
        [
            ["[Thành viên 1]", "[MSSV]", "Frontend, trải nghiệm học tập, biểu đồ", "...%", "Bổ sung chữ ký"],
            ["[Thành viên 2]", "[MSSV]", "Backend, dữ liệu, thanh toán, bảo mật", "...%", "Bổ sung chữ ký"],
            ["[Thành viên 3]", "[MSSV]", "Judge0, AI local, kiểm thử, tài liệu", "...%", "Bổ sung chữ ký"],
        ],
    )


def add_problem_section(doc: Document, assets: dict[str, Path], retention_formula: Path, competition_formula: Path) -> None:
    add_heading(doc, "1. ĐẶT VẤN ĐỀ", 1, page_break=True)
    add_heading(doc, "1.1. Mô tả bài toán", 2)
    add_paragraph(
        doc,
        "Việc học lập trình trực tuyến thường bị phân mảnh: học liệu nằm ở một nơi, công cụ chạy code ở nơi khác, chatbot không hiểu nội dung khóa học và quản trị viên chỉ nhìn thấy phần trăm tiến độ sau khi học viên đã bỏ học. Với người mới, khoảng cách giữa đọc lý thuyết và nhận phản hồi đúng lúc làm tăng khả năng mất nhịp; với đơn vị đào tạo, việc phụ thuộc quota AI hoặc dịch vụ chấm code bên ngoài làm giảm khả năng chủ động vận hành.",
    )
    add_paragraph(
        doc,
        "CodeLearn được xây dựng cho các khóa SQL, C, C++ và Python, đồng thời để mở khả năng bổ sung Java, JavaScript, Go, Rust và các runtime khác. Hệ thống tích hợp bài học, ví dụ chạy trực tiếp, bài tập có test ẩn, quiz, thanh toán, tiến độ, AI Tutor có nguồn, thi đua có thưởng, Trạm giữ nhịp và dashboard Can thiệp sớm.",
    )
    add_paragraph(
        doc,
        "Bài toán trọng tâm của đề tài không chỉ là cung cấp nội dung. Nhóm em cần xây dựng một vòng kín có thể giải thích: thu thập hành vi học thật, tính điểm giữ nhịp, cho người học thấy xu hướng tăng hoặc giảm, giao nhiệm vụ phù hợp, gửi thông báo và giúp quản trị viên can thiệp trước khi học viên trả phí rời khỏi lộ trình.",
    )

    add_heading(doc, "1.2. Mục đích", 2)
    add_paragraph(doc, "Tài liệu này mô tả cách nhóm em hiện thực các yêu cầu của CodeLearn và làm cơ sở cho triển khai, kiểm thử, bảo trì và mở rộng hệ thống.")
    add_bullets(
        doc,
        [
            "Mô tả kiến trúc frontend, backend, cơ sở dữ liệu và các local service.",
            "Giải thích các quyết định thiết kế đối với Judge0, Llama local, RAG, thanh toán và bảo mật.",
            "Đặc tả quy trình nghiệp vụ, actor, use case, sequence và mô hình dữ liệu.",
            "Trình bày công thức giữ nhịp và thi đua theo phiên bản, có giới hạn chống cày điểm.",
            "Đưa ra bằng chứng kiểm thử và giới hạn quy mô để tránh cam kết vượt dữ liệu thực nghiệm.",
        ],
    )

    add_heading(doc, "1.3. Quy ước tài liệu", 2)
    add_table(
        doc,
        "Thuật ngữ và quy ước sử dụng",
        ["Thuật ngữ", "Diễn giải"],
        [
            ["PAID / PENDING", "Trạng thái đã thanh toán / đang chờ của CoursePurchase."],
            ["Submission", "Một lần nộp code để chấm bằng test case server-side."],
            ["RAG", "Truy xuất bài học liên quan trước khi tạo ngữ cảnh cho LLM."],
            ["Health score", "Điểm giữ nhịp 0-100; không phải chẩn đoán tâm lý hay dự đoán ML."],
            ["p95", "95% yêu cầu có thời gian phản hồi nhỏ hơn hoặc bằng giá trị này."],
            ["N/A", "Chưa có dữ liệu hoặc không áp dụng; không được thay bằng kết quả suy đoán."],
            ["Tiền tệ", "Giá trị hiển thị theo định dạng Việt Nam, ví dụ 2.000 đồng."],
        ],
    )

    add_heading(doc, "1.4. Các yêu cầu nghiệp vụ", 2)
    requirements = [
        ("1.4.1. Quản lý tài khoản và quyền truy cập", [
            "Đăng ký, xác minh email, đăng nhập, đăng xuất và khôi phục mật khẩu.",
            "Phân quyền LEARNER/ADMIN ở backend; nội dung đã mua chỉ mở khi purchase là PAID.",
        ]),
        ("1.4.2. Học liệu, thực hành và đánh giá", [
            "Quản lý khóa học, bài học, ví dụ, bài tập, test case và quiz.",
            "Chạy code, chấm test ẩn, lưu submission và cập nhật tiến độ khi đạt.",
        ]),
        ("1.4.3. Thanh toán và thông báo", [
            "Tạo VietQR với mã duy nhất; nhận webhook SePay có API key và xử lý idempotent.",
            "Mở khóa đúng một lần, tạo thông báo trong ứng dụng và email tùy cấu hình.",
        ]),
        ("1.4.4. Thi đua, phòng thi và phần thưởng", [
            "Admin cấu hình mùa, đề, thời gian, thang thưởng; học viên start/submit attempt.",
            "Ranking lấy dữ liệu trong mùa, chuẩn hóa 1.000 điểm và chống cộng lặp.",
        ]),
        ("1.4.5. Giữ nhịp và can thiệp sớm", [
            "Chỉ đánh giá học viên đã thanh toán; trả điểm thành phần, lý do và xu hướng 28 ngày.",
            "Sinh nhiệm vụ hôm nay, gói cứu nhịp 48 giờ; admin theo dõi và giao can thiệp.",
        ]),
        ("1.4.6. Quản trị và khả năng mở rộng nội dung", [
            "CRUD khóa, bài, bài tập, đơn hàng, contest, reward claim và dashboard rủi ro.",
            "Ngôn ngữ là dữ liệu chuỗi, runtime được phát hiện từ Judge0 thay vì enum đóng.",
        ]),
    ]
    for heading, items in requirements:
        add_heading(doc, heading, 3)
        add_bullets(doc, items)

    add_heading(doc, "1.5. Yêu cầu phi nghiệp vụ", 2)
    nfrs = [
        ("1.5.1. Hiệu năng (Performance)", "API thông thường hướng đến p95 dưới 500 ms trên server mục tiêu. Code ngắn hướng đến p95 dưới 5 giây khi hàng đợi bình thường. Bảng xếp hạng dùng cache ngắn; truy vấn retention có index theo user, status và thời gian."),
        ("1.5.2. Tính sẵn sàng (Availability)", "Container có restart policy; Judge0 API, PostgreSQL và Redis có health check. AI hoặc email lỗi không làm mất chức năng học tập cốt lõi."),
        ("1.5.3. Bảo mật (Security)", "Mật khẩu băm Argon2id; JWT trong cookie HttpOnly; Helmet, CORS, Zod, role middleware, rate limit và test ẩn ở backend. Code bị giới hạn CPU/RAM/thời gian và tắt mạng."),
        ("1.5.4. Khả năng mở rộng (Scalability)", "Mục tiêu khởi điểm là 1.000 tài khoản và 50-100 người hoạt động đồng thời, không đồng nghĩa 1.000 lượt biên dịch cùng lúc. Judge0 worker có thể scale ngang; production nhiều API instance cần Redis cho queue/cache/limiter dùng chung."),
        ("1.5.5. Tính bảo trì (Maintainability)", "TypeScript, module theo miền, schema validation, Prisma migration và công thức có version. Thay đổi trọng số phải có test và không ghi đè lịch sử diễn giải."),
        ("1.5.6. Tính tương thích đa nền tảng (Compatibility)", "Frontend responsive cho desktop/mobile; backend và local service chạy qua Node.js/Docker trên môi trường hỗ trợ. Judge0 cần Linux/cgroup phù hợp ở production."),
        ("1.5.7. Tính dễ sử dụng (Usability)", "Trạng thái loading/error/empty rõ ràng, editor giữ code sau lỗi, biểu đồ có nhãn vùng rủi ro và nhiệm vụ có CTA trực tiếp."),
        ("1.5.8. Sao lưu và phục hồi dữ liệu", "Production phải backup PostgreSQL hằng ngày, kiểm tra khôi phục, lưu migration và tách volume Judge0/Redis. Máy demo chưa thay thế quy trình backup production."),
        ("1.5.9. Ghi log và kiểm toán", "Docker log rotation 10 MB x 3 file; webhook lưu event/transaction. Hướng tiếp theo là audit log admin và dashboard CPU, RAM, queue depth, p95 và tỷ lệ lỗi."),
        ("1.5.10. Tuân thủ pháp lý", "Chỉ thu thập dữ liệu cần cho học tập; không công bố điểm giữ nhịp như kết luận tâm lý; bí mật thanh toán/email không đưa vào repository hoặc báo cáo."),
    ]
    for heading, text in nfrs:
        add_heading(doc, heading, 3)
        add_paragraph(doc, text)

    add_heading(doc, "1.6. Các kỹ thuật áp dụng để giải quyết bài toán", 2)
    add_table(
        doc,
        "Công nghệ và lý do lựa chọn",
        ["Lớp", "Công nghệ", "Vai trò và lý do"],
        [
            ["Frontend", "React 18, TypeScript, Vite", "SPA nhiều tương tác, component tái sử dụng, kiểm tra kiểu."],
            ["Giao diện", "Tailwind CSS, Lucide, CodeMirror", "Responsive, icon chuẩn và editor code theo ngôn ngữ."],
            ["Backend", "Node.js, Express, TypeScript, Zod", "REST API module hóa, validation tại biên."],
            ["Dữ liệu", "PostgreSQL, Prisma", "Quan hệ, transaction, migration và ràng buộc unique/index."],
            ["Chấm code", "Judge0 CE, Redis, PostgreSQL", "Sandbox, test nhiều ngôn ngữ, worker và hàng đợi tự chủ."],
            ["AI", "llama.cpp, Qwen GGUF, RAG", "OpenAI-compatible API local, không phụ thuộc quota cloud."],
            ["Thanh toán", "VietQR, SePay webhook/IPN", "Đối soát chuyển khoản và mở khóa tự động."],
            ["Kiểm thử", "Vitest, Supertest, Playwright", "Unit, integration, E2E desktop/mobile và smoke test."],
        ],
    )
    add_figure(doc, retention_formula, "Trọng số công thức điểm giữ nhịp V3", 15.5)
    add_figure(doc, competition_formula, "Thang điểm thi đua V2 tối đa 1.000", 15.5)


def add_business_section(doc: Document, assets: dict[str, Path]) -> None:
    add_heading(doc, "2. PHÂN TÍCH NGHIỆP VỤ", 1, page_break=True)
    add_heading(doc, "2.1. Các lớp người dùng hệ thống", 2)
    add_table(
        doc,
        "Lớp người dùng và quyền chính",
        ["Lớp người dùng", "Đặc điểm", "Quyền chính"],
        [
            ["Khách", "Chưa có phiên hoặc chưa mua khóa", "Xem khóa, đăng ký, đăng nhập, tìm kiếm và checkout."],
            ["Học viên", "Tài khoản LEARNER; quyền học theo purchase", "Học, chạy/nộp code, quiz, tiến độ, AI, contest, retention, thông báo."],
            ["Quản trị viên", "Tài khoản ADMIN", "CRUD học liệu, đơn hàng, contest, reward, retention risk và intervention."],
            ["Dịch vụ tích hợp", "SePay, SMTP/EmailJS, Judge0, llama.cpp", "Gửi sự kiện hoặc cung cấp năng lực qua API được kiểm soát."],
        ],
    )

    add_heading(doc, "2.2. Các đối tác nghiệp vụ và thừa tác viên", 2)
    add_table(
        doc,
        "Stakeholder và tiêu chí thành công",
        ["Bên liên quan", "Nhu cầu", "Tiêu chí thành công"],
        [
            ["Học viên chưa mua", "Hiểu khóa và giá", "Không bị gắn nhãn bỏ học; thanh toán rõ trạng thái."],
            ["Học viên đang học", "Phản hồi nhanh và lộ trình rõ", "Submission/quiz/progress nhất quán; biết mình tăng hay giảm nhịp."],
            ["Học viên thi đua", "Điểm công bằng", "Điểm có breakdown; phần thưởng và deadline minh bạch."],
            ["Quản trị viên", "Quản lý và can thiệp", "Nhìn đúng nhóm trả phí có nguy cơ; thao tác đúng quyền."],
            ["Giảng viên", "Đánh giá tính đúng và độ khó", "Có sơ đồ, công thức, mã nguồn, test và giới hạn rõ."],
        ],
    )

    add_heading(doc, "2.3. Các quy trình nghiệp vụ", 2)
    add_heading(doc, "2.3.1. Quy trình dành cho khách và học viên", 3)
    add_paragraph(doc, "Các quy trình chính được thiết kế liên tục từ lúc tìm khóa đến khi quay lại học sau can thiệp.")
    add_table(
        doc,
        "Quy trình đăng ký và thanh toán",
        ["Bước", "Thao tác", "Dữ liệu/điều kiện", "Ngoại lệ"],
        [
            ["1", "Đăng ký và xác minh", "Email, tên, mật khẩu; token còn hạn", "Email trùng hoặc token sai/hết hạn"],
            ["2", "Chọn khóa", "Course đang tồn tại, giá 2.000 đồng", "Khóa không tồn tại"],
            ["3", "Tạo purchase và QR", "paymentCode duy nhất, status PENDING", "Purchase cũ PAID"],
            ["4", "SePay gửi webhook", "API key đúng, tiền vào, đủ số tiền", "Sai key, thiếu tiền, event trùng"],
            ["5", "Mở khóa", "Transaction + purchase PAID trong transaction", "Rollback nếu cập nhật thất bại"],
            ["6", "Thông báo", "In-app luôn có; email theo preference", "Email lỗi không rollback purchase"],
        ],
    )
    add_table(
        doc,
        "Quy trình học, thực hành và giữ nhịp",
        ["Giai đoạn", "Luồng chính", "Dữ liệu được tạo"],
        [
            ["Học bài", "Mở lesson, đọc nội dung, chạy example", "Progress, streak, lastActiveDate"],
            ["Làm bài", "Nộp code, Judge0 chạy test, so khớp output", "Submission, error fingerprint, progress khi pass"],
            ["Làm quiz", "Gửi lựa chọn, tính score/total", "QuizAttempt và progress"],
            ["Xem giữ nhịp", "Tính V3 và dựng 28 mốc", "Breakdown, reasons, trend, missions"],
            ["Can thiệp", "Nhận gói 48 giờ và hoàn thành nhiệm vụ", "Intervention, mission status, outcome"],
            ["Thi đua", "Hoạt động trong mùa và phòng thi", "Breakdown /1.000, rank, reward claim"],
        ],
    )

    add_heading(doc, "2.3.2. Quy trình dành cho quản trị viên", 3)
    add_table(
        doc,
        "Quy trình quản trị và can thiệp",
        ["Quy trình", "Đầu vào", "Xử lý", "Đầu ra"],
        [
            ["Quản lý học liệu", "Course, lesson, exercise, quiz", "Validate, CRUD, sắp thứ tự", "Nội dung mới hoặc cập nhật"],
            ["Quản lý đơn", "Purchase và bằng chứng", "Xem trạng thái, xác nhận theo quyền", "PAID/PENDING và notification"],
            ["Cấu hình mùa", "Thời gian, đề, phần thưởng", "Validate rank range và problem", "Contest có room/ranking"],
            ["Duyệt thưởng", "RewardClaim PENDING", "Approve/reject, ghi note", "Claim mới và thông báo"],
            ["Can thiệp sớm", "Danh sách learner PAID", "Sắp rủi ro V3, xem weakest course", "Gói 48 giờ và lịch sử outcome"],
        ],
    )

    add_heading(doc, "2.3.3. Quy trình của dịch vụ tích hợp", 3)
    add_bullets(
        doc,
        [
            "Judge0 nhận submission bất đồng bộ, trả token, worker lấy job từ Redis và backend poll kết quả đến khi hoàn tất hoặc timeout.",
            "llama.cpp nhận messages theo OpenAI-compatible API; backend thực hiện guardrail và RAG trước khi gọi model.",
            "SePay gọi webhook công khai; backend xác thực API key, kiểm tra giao dịch và xử lý idempotent.",
            "SMTP/EmailJS chỉ là kênh bổ sung; Notification trong PostgreSQL là nguồn trạng thái chính trong ứng dụng.",
        ],
    )


def add_requirement_section(doc: Document, assets: dict[str, Path]) -> None:
    add_heading(doc, "3. PHÂN TÍCH YÊU CẦU", 1, page_break=True)
    add_heading(doc, "3.1. Xác định yêu cầu các bên liên quan", 2)
    stakeholder_groups = [
        ("3.1.1. Học viên", ["Nội dung đúng quyền mua; editor dễ dùng; kết quả chấm không lộ test ẩn.", "Điểm giữ nhịp và thi đua phải giải thích được, không cộng lặp vô hạn.", "Thanh toán và thông báo phải rõ trạng thái, không yêu cầu thao tác kỹ thuật."]),
        ("3.1.2. Quản trị viên", ["Quản lý đầy đủ học liệu, đơn hàng, thi đua và phần thưởng.", "Danh sách can thiệp chỉ gồm học viên PAID, sắp theo rủi ro thật và có hành động gợi ý.", "Không thể truy cập route admin bằng tài khoản học viên dù sửa URL frontend."]),
        ("3.1.3. Giảng viên và đơn vị vận hành", ["Có tài liệu kiến trúc, use case, sequence, ERD và ma trận truy vết.", "Có test tự động, smoke test, load test và hướng dẫn chạy local.", "Không tuyên bố 1.000 biên dịch đồng thời khi chưa đo trên server thật."]),
    ]
    for heading, bullets in stakeholder_groups:
        add_heading(doc, heading, 3)
        add_bullets(doc, bullets)

    add_heading(doc, "3.2. Xác định actor", 2)
    add_table(
        doc,
        "Actor và phạm vi tương tác",
        ["Actor", "Loại", "Use case tiêu biểu"],
        [
            ["Khách", "Người dùng", "Đăng ký, đăng nhập, xem khóa, checkout"],
            ["Học viên", "Người dùng", "Học, nộp bài, quiz, AI, contest, retention, notification"],
            ["Quản trị viên", "Người dùng", "CRUD, duyệt đơn/thưởng, cấu hình contest, can thiệp"],
            ["SePay", "Hệ thống ngoài", "Gửi webhook/IPN xác nhận giao dịch"],
            ["Judge0", "Dịch vụ nội bộ", "Biên dịch/chạy code trong sandbox"],
            ["Llama local", "Dịch vụ nội bộ", "Sinh câu trả lời theo prompt/RAG"],
            ["Email provider", "Hệ thống ngoài", "Gửi email theo preference"],
        ],
    )

    add_heading(doc, "3.3. Use Case Diagram", 2)
    add_heading(doc, "3.3.1. Tổng quan hệ thống", 3)
    add_figure(doc, assets["use-case.png"], "Use Case Diagram tổng quan CodeLearn", 15.8)

    add_heading(doc, "3.3.2. Phân rã tác nhân Học viên", 3)
    add_table(
        doc,
        "Nhóm use case của Học viên",
        ["Nhóm", "Use case", "Ràng buộc"],
        [
            ["Tài khoản", "Đăng ký, đăng nhập, reset password", "Email/token/password hợp lệ"],
            ["Học tập", "Mở lesson, chạy example, note/bookmark", "Purchase PAID với nội dung khóa"],
            ["Đánh giá", "Nộp exercise, làm quiz", "Không trả test ẩn; cập nhật progress khi đạt"],
            ["Hỗ trợ", "Hỏi AI Tutor, xem Code DNA", "Guardrail; dữ liệu submission của chính mình"],
            ["Động lực", "Retention, contest, reward claim", "Chỉ dữ liệu trong phạm vi và thời gian hợp lệ"],
        ],
    )

    add_heading(doc, "3.3.3. Phân rã tác nhân Quản trị viên", 3)
    add_table(
        doc,
        "Nhóm use case của Quản trị viên",
        ["Nhóm", "Use case", "Kiểm soát"],
        [
            ["Học liệu", "CRUD course/lesson/exercise", "Zod + role ADMIN"],
            ["Người dùng/đơn", "Xem user, purchase, xác nhận đơn", "Không lộ passwordHash; lưu trạng thái"],
            ["Thi đua", "CRUD contest/problem/reward, duyệt claim", "Validate thời gian, rank range và liên kết"],
            ["Giữ nhịp", "Xem risk, giao intervention", "Chỉ learner PAID; dùng cùng công thức V3"],
        ],
    )

    add_heading(doc, "3.3.4. Phân rã tác nhân Dịch vụ tích hợp", 3)
    add_bullets(doc, [
        "SePay chỉ được gọi endpoint webhook/IPN và phải vượt xác thực tương ứng.",
        "Judge0 chỉ nằm trong private/local network ở production; code người học không bật network.",
        "Llama local không truy cập database trực tiếp; chỉ nhận context đã giới hạn từ backend.",
        "Email provider không quyết định trạng thái nghiệp vụ; lỗi gửi email không rollback dữ liệu chính.",
    ])

    add_heading(doc, "3.3.5. Đặc tả các ca sử dụng", 3)
    use_cases = [
        ("UC-01 - Nộp bài lập trình", [
            ["Actor", "Học viên hoặc khách ở phạm vi cho phép"],
            ["Tiền điều kiện", "Exercise tồn tại; Judge0 có runtime; payload hợp lệ"],
            ["Luồng chính", "Nhập code → submit → chạy từng test → so khớp → lưu kết quả → cập nhật progress nếu pass"],
            ["Luồng phụ", "Compile error, timeout, queue đầy, runtime không có, Judge0 mất kết nối"],
            ["Hậu điều kiện", "Submission có status/passedCount/totalCount; test ẩn không lộ"],
        ]),
        ("UC-02 - Xác nhận thanh toán", [
            ["Actor", "Học viên, SePay, hệ thống"],
            ["Tiền điều kiện", "Purchase PENDING và paymentCode còn hiệu lực"],
            ["Luồng chính", "Tạo QR → chuyển khoản → webhook → xác thực → transaction → PAID → notification"],
            ["Luồng phụ", "Sai key, tiền ra, thiếu mã, thiếu tiền, event/transaction trùng"],
            ["Hậu điều kiện", "Purchase chỉ PAID một lần; event được lưu để đối soát"],
        ]),
        ("UC-03 - Can thiệp sớm", [
            ["Actor", "Admin và học viên PAID"],
            ["Tiền điều kiện", "Có dữ liệu học và tài khoản đã thanh toán"],
            ["Luồng chính", "Tính V3 → phân nhóm → xem lý do/trend → giao gói → notification → đồng bộ nhiệm vụ"],
            ["Luồng phụ", "Người chưa mua nhận NOT_STARTED; email lỗi nhưng in-app còn"],
            ["Hậu điều kiện", "Intervention có baseline, deadline, completedMissions, status và outcome"],
        ]),
        ("UC-04 - Tham gia phòng thi", [
            ["Actor", "Học viên đủ điều kiện"],
            ["Tiền điều kiện", "Contest ACTIVE, có problem, learner có purchase phù hợp"],
            ["Luồng chính", "Start attempt → làm exercise/quiz → submit → chuẩn hóa điểm → cập nhật ranking"],
            ["Luồng phụ", "Hết giờ, contest chưa mở, không có đề, attempt không tồn tại"],
            ["Hậu điều kiện", "Attempt SUBMITTED/EXPIRED với score/maxScore; leaderboard cache được làm mới"],
        ]),
    ]
    for title, rows in use_cases:
        add_table(doc, title, ["Mục", "Nội dung"], rows)

    add_heading(doc, "3.3.6. Biểu đồ hoạt động", 3)
    add_paragraph(doc, "Các biểu đồ dưới thể hiện thứ tự tương tác và nhánh lỗi của ba quy trình có rủi ro kỹ thuật cao nhất.")
    add_figure(doc, assets["submit-sequence.png"], "Luồng nộp bài và chấm test ẩn", 15.5)
    add_figure(doc, assets["payment-sequence.png"], "Luồng thanh toán VietQR và SePay webhook", 15.5)
    add_figure(doc, assets["retention-loop.png"], "Vòng kín giữ nhịp và can thiệp sớm", 15.5)


def add_design_section(doc: Document, assets: dict[str, Path], retention_screenshot: Path) -> None:
    add_heading(doc, "4. THIẾT KẾ HỆ THỐNG", 1, page_break=True)
    add_heading(doc, "4.1. Kiến trúc", 2)
    add_paragraph(
        doc,
        "CodeLearn dùng kiến trúc client-server nhiều lớp. React SPA chịu trách nhiệm hiển thị và tương tác; Express tổ chức nghiệp vụ theo module; Prisma truy cập PostgreSQL; Judge0 và llama.cpp là dịch vụ nội bộ qua API. SePay/email là tích hợp ngoài nhưng không thay thế nguồn dữ liệu nghiệp vụ trong hệ thống.",
    )
    add_figure(doc, assets["architecture.png"], "Kiến trúc tổng thể CodeLearn", 15.8)
    add_table(
        doc,
        "Module backend và trách nhiệm",
        ["Module", "Endpoint tiêu biểu", "Trách nhiệm"],
        [
            ["auth", "/api/auth/login, /me", "Xác thực, token, OAuth, password reset"],
            ["content/admin", "/api/courses, /api/admin/...", "Học liệu và CRUD theo quyền"],
            ["run/exercise", "/api/run, /exercises/:id/submit", "Runtime discovery, queue, chấm test"],
            ["quiz/progress", "/api/quizzes/:id/submit, /progress", "Đánh giá và hoàn thành"],
            ["ai", "/api/ai/chat/stream", "Quick diagnostic, guardrail, RAG và LLM"],
            ["payment", "/api/payments/sepay/webhook", "Purchase, event, transaction, mở khóa"],
            ["contest", "/api/contests", "Room, score /1.000, rank và reward"],
            ["retention", "/api/retention/plan", "V3, trend, mission, intervention"],
            ["notification", "/api/notifications", "In-app, preference và email bổ sung"],
        ],
    )

    add_heading(doc, "4.1.1. Kiến trúc triển khai local", 3)
    add_figure(doc, assets["local-deployment.png"], "Triển khai local và luồng webhook công khai", 15.8)
    add_table(
        doc,
        "Cổng và biến cấu hình chính",
        ["Thành phần", "Địa chỉ", "Biến cấu hình", "Ghi chú"],
        [
            ["Frontend", "localhost:5173", "CLIENT_ORIGIN", "Vite dev hoặc static build"],
            ["Backend", "localhost:4000", "PORT, PUBLIC_API_URL", "REST API"],
            ["Judge0", "localhost:2358", "JUDGE0_URL", "CE 1.13.1, private ở production"],
            ["Llama", "localhost:8080/v1", "LOCAL_LLM_BASE_URL", "OpenAI-compatible"],
            ["PostgreSQL", "Theo DATABASE_URL", "DATABASE_URL", "Không đưa secret vào tài liệu"],
            ["Webhook", "Quick Tunnel khi demo", "PUBLIC_API_URL", "Production dùng domain cố định"],
        ],
    )

    add_heading(doc, "4.1.2. Kiến trúc mở rộng cho 1.000 học viên", 3)
    add_paragraph(
        doc,
        "Mục tiêu 1.000 học viên được hiểu là 1.000 tài khoản và khoảng 50-100 người hoạt động đồng thời ở giai đoạn đầu. Nếu 1.000 người cùng submit code, cấu hình hiện tại không chạy 1.000 sandbox song song: backend chỉ nhận theo giới hạn, đưa vào queue và trả 429 khi đầy để bảo vệ hệ thống.",
    )
    add_table(
        doc,
        "Cấu hình khởi điểm trên server mục tiêu",
        ["Thành phần", "Cấu hình đề xuất", "Cách mở rộng"],
        [
            ["CodeLearn API", "2 instance, 2 vCPU, 2-4 GB/instance", "Load balancer; Redis limiter/cache"],
            ["PostgreSQL app", "2-4 vCPU, 4-8 GB, SSD", "Managed DB/read replica khi cần"],
            ["Judge0 API", "1-2 instance", "Private network"],
            ["Judge0 worker", "2-4 container, 2 vCPU, 1,5-2 GB", "Scale ngang theo queue depth"],
            ["Redis/PostgreSQL Judge0", "Tách volume, backup", "Theo dõi queue, disk và connection"],
            ["Llama/RAG", "Máy riêng; ưu tiên GPU", "Batching, parallel slots, model phù hợp"],
        ],
    )

    add_heading(doc, "4.2. Class và Sequence Diagram", 2)
    add_heading(doc, "4.2.1. Chạy và chấm code", 3)
    add_table(
        doc,
        "Các lớp/module trong luồng chấm code",
        ["Thành phần", "Trách nhiệm", "Quan hệ"],
        [
            ["RunController", "Validate language/source/stdin", "Gọi CodeRunner"],
            ["CodeRunner", "Chọn runtime, semaphore, submit/poll", "Gọi Judge0 API"],
            ["ExerciseService", "Chạy test, so output, lưu kết quả", "Submission + Progress"],
            ["Judge0", "Sandbox compile/run", "Redis queue + worker"],
            ["Submission", "Lưu code, status, counts, error DNA", "Thuộc User và Exercise"],
        ],
    )
    add_figure(doc, assets["submit-sequence.png"], "Sequence chạy code và chấm test ẩn", 15.6)

    add_heading(doc, "4.2.2. AI Tutor local", 3)
    add_paragraph(doc, "Quick reply xử lý lỗi phổ biến trước để giảm độ trễ. Với câu hỏi cần kiến thức, backend kiểm tra phạm vi, truy xuất tối đa ngữ cảnh liên quan rồi stream token từ llama.cpp.")
    add_figure(doc, assets["ai-sequence.png"], "Sequence AI Tutor local kết hợp RAG và guardrail", 15.6)

    add_heading(doc, "4.2.3. Thanh toán", 3)
    add_figure(doc, assets["payment-sequence.png"], "Sequence thanh toán và mở khóa idempotent", 15.6)

    add_heading(doc, "4.2.4. Thi đua và phần thưởng", 3)
    add_paragraph(
        doc,
        "COMPETITION_V2_1000 gồm học bài 200, bài code 250, quiz 150, phòng thi 300 và học đều 100. Lesson/exercise/quiz chỉ lấy đối tượng duy nhất hoặc kết quả tốt nhất trong thời gian mùa. Khi bằng điểm, ưu tiên phòng thi, thực hành, quiz, độ đều và thời điểm đạt sớm hơn.",
    )
    add_figure(doc, assets["contest-state.png"], "Trạng thái contest, attempt và reward claim", 15.4)

    add_heading(doc, "4.2.5. Giữ nhịp và can thiệp", 3)
    add_paragraph(
        doc,
        "Điểm RETENTION_V3_2026_07 là tổng năm nhóm: gần nhất 25, độ đều 20, tiến độ 20, chất lượng 25 và xu hướng 10. Tiến độ tổng chỉ đóng vai trò nền; phần lớn điểm đến từ hành vi gần đây. Nghỉ từ 14 ngày luôn AT_RISK; nghỉ từ 7 ngày không thể ON_TRACK.",
    )
    add_code(
        doc,
        "health = recency(25) + consistency(20) + progress(20)\n"
        "       + mastery(25) + momentum(10)\n\n"
        "consistency = min(round(activeDays14 / 6 * 20), 20)\n"
        "progress = round(sqrt(overallPercent / 100) * 8)\n"
        "           + min(round(completedItems14 / 4 * 12), 12)\n"
        "mastery = passRatioUnique * 15 + quizBestAverage * 10\n"
        "risk: >=70 ON_TRACK; >=45 WATCH; <45 AT_RISK\n"
        "hard gates: inactive>=14 => AT_RISK; inactive>=7 => not ON_TRACK",
    )
    add_figure(doc, assets["retention-loop.png"], "Sequence vòng kín giữ nhịp", 15.4)

    add_heading(doc, "4.3. Thiết kế giao diện", 2)
    add_paragraph(doc, "Giao diện dùng layout responsive, trạng thái loading/error/empty và CTA theo luồng chính. Các ảnh dưới là bề mặt quan trọng dùng khi kiểm thử và bảo vệ.")
    add_figure(doc, assets["courses-ui.png"], "Danh sách khóa học trên giao diện học viên", 14.8)
    add_figure(doc, assets["lesson-ui.png"], "Bài học kết hợp nội dung và ví dụ thực hành", 14.8)
    add_figure(doc, assets["exercise-ui.png"], "Bài tập code với CodeMirror và test mẫu", 14.8)
    add_figure(doc, retention_screenshot, "Trạm giữ nhịp với biểu đồ xu hướng 28 ngày", 15.2)
    add_figure(doc, assets["admin-retention-ui.png"], "Dashboard Can thiệp sớm cho quản trị viên", 15.0)
    add_figure(doc, assets["contest-ui.png"], "Trang mùa thi, ranking và ưu đãi", 15.0)
    add_figure(doc, assets["admin-contest-ui.png"], "Giao diện quản trị mùa thi và phòng thi", 15.0)
    add_figure(doc, assets["notifications-ui.png"], "Trung tâm thông báo và tùy chọn email", 15.0)
    add_figure(doc, assets["code-dna-ui.png"], "Hồ sơ lỗi Code DNA", 15.0)

    add_heading(doc, "4.4. Thiết kế cơ sở dữ liệu", 2)
    add_paragraph(
        doc,
        "Prisma schema hiện có 30 model và 13 enum. Ngôn ngữ lập trình được lưu VARCHAR(80), cho phép thêm runtime mà không phải sửa enum database. Hai migration mới chuyển kiểu language và bổ sung index cho truy vấn retention/ranking; toàn dự án có 18 migration.",
    )
    add_figure(doc, assets["erd.png"], "ERD mức miền nghiệp vụ của CodeLearn", 15.8)
    add_table(
        doc,
        "Mô hình dữ liệu trọng tâm",
        ["Miền", "Model chính", "Ràng buộc quan trọng"],
        [
            ["Tài khoản", "User, RegistrationVerification, PasswordResetToken", "email unique; token hash; role"],
            ["Học liệu", "Course, Lesson, Example, Exercise, TestCase, Quiz", "slug/order; hidden test; language VARCHAR"],
            ["Học tập", "Progress, Submission, QuizAttempt, UserBadge", "unique progress item; index user/time/status"],
            ["Thanh toán", "CoursePurchase, SepayWebhookEvent, PaymentTransaction", "unique user-course, paymentCode, event/transaction"],
            ["Thi đua", "Contest, ContestProblem, ContestAttempt, ContestReward, RewardClaim", "rank range; attempt status; unique claim"],
            ["Giữ nhịp", "LearningIntervention, InterventionMission", "baseline, deadline, status, outcome"],
            ["Thông báo", "Notification, NotificationPreference", "dedupeKey theo user"],
        ],
    )
    add_table(
        doc,
        "Index phục vụ dữ liệu tăng trưởng",
        ["Bảng", "Index", "Truy vấn được hỗ trợ"],
        [
            ["Submission", "userId, status, createdAt", "Pass duy nhất và cửa sổ 14/30/56 ngày"],
            ["QuizAttempt", "userId, createdAt", "Quiz gần đây và kết quả tốt nhất"],
            ["Progress", "userId, completed, completedAt", "Tiến độ và trend lịch sử"],
            ["ContestAttempt", "contestId, status, userId", "Leaderboard theo mùa"],
            ["CoursePurchase", "status, userId", "Lọc đúng học viên PAID"],
        ],
    )


def add_appendices(doc: Document) -> None:
    add_heading(doc, "PHỤ LỤC A. HIỆN THỰC VÀ KIỂM THỬ", 1, page_break=True)
    add_heading(doc, "A.1. Cấu trúc dự án", 2)
    add_code(
        doc,
        "client/src/                 React pages, components, API adapters\n"
        "server/src/modules/         Auth, content, run, exercise, AI, payment, contest, retention\n"
        "server/prisma/              Schema, migrations, seed data\n"
        "docker/local-services/      Judge0, PostgreSQL, Redis, llama.cpp\n"
        "scripts/                    Demo, smoke test, load test, report generation\n"
        "e2e/                        Playwright critical flows\n"
        "docs/                       Báo cáo và tài liệu vận hành",
    )

    add_heading(doc, "A.2. Kết quả kiểm thử cập nhật", 2)
    add_table(
        doc,
        "Tổng hợp kiểm thử ngày 28/07/2026",
        ["Hạng mục", "Quy mô", "Kết quả", "Ghi chú"],
        [
            ["Server unit/integration", "17 tệp / 76 test", "76/76 đạt", "Auth, payment, AI, retention, contest, runner..."],
            ["Client unit", "1 tệp / 4 test", "4/4 đạt", "Định dạng và dữ liệu giao diện"],
            ["E2E desktop/mobile", "3 kịch bản x 2 project", "6/6 đạt", "Edge desktop và Pixel 7 emulation"],
            ["Build", "Server + client", "Đạt", "TypeScript và Vite production build"],
            ["Lint", "Client + server", "Đạt", "ESLint không có lỗi"],
            ["Smoke", "4 service + API nghiệp vụ", "Đạt", "Frontend, backend, Judge0, Llama, webhook auth"],
            ["Tổng test tự động", "80 unit/integration + 6 E2E", "86/86 đạt", "Không đồng nhất với code coverage"],
        ],
    )
    add_table(
        doc,
        "Kết quả load smoke Judge0 trên máy demo",
        ["Tổng lượt", "Đồng thời", "Kết quả", "p95", "Throughput"],
        [
            ["20", "4", "20/20", "2,38 giây", "2,80 lượt/giây"],
            ["50", "10", "50/50", "3,91 giây", "3,71 lượt/giây"],
            ["100", "20", "100/100", "6,24 giây", "3,78 lượt/giây"],
        ],
    )
    add_paragraph(
        doc,
        "Load test dùng chương trình Python rất ngắn với hai container worker, mỗi container COUNT=2. C, C++, Python và SQL đã được kiểm tra riêng và đều Accepted. Số liệu này chứng minh cơ chế hàng đợi ở quy mô thử nghiệm, không chứng minh 1.000 bài phức tạp chạy đồng thời.",
    )

    add_heading(doc, "A.3. Đánh giá giới hạn", 2)
    add_table(
        doc,
        "Giới hạn và hướng xử lý",
        ["Giới hạn", "Ảnh hưởng", "Hướng xử lý"],
        [
            ["Qwen 1.5B local", "Suy luận dài chưa mạnh", "Model lớn/GPU, eval theo bộ câu hỏi khóa học"],
            ["1.000 submit cùng lúc", "Queue đầy và trả 429", "JobId bất đồng bộ, Redis queue toàn cục, scale worker"],
            ["Retention là heuristic", "Chưa chứng minh uplift", "Thu thập 4-8 tuần, cohort/control và hiệu chỉnh trọng số"],
            ["Cache trong process", "Không đồng bộ nhiều API instance", "Chuyển cache/limiter sang Redis"],
            ["Quick Tunnel", "URL đổi, không phù hợp production", "Domain cố định, HTTPS, reverse proxy"],
            ["Bundle editor lớn", "Tải chậm mạng yếu", "Code splitting và lazy load extension"],
        ],
    )

    add_heading(doc, "PHỤ LỤC B. HƯỚNG DẪN CHẠY", 1, page_break=True)
    add_heading(doc, "B.1. Khởi động local service", 2)
    add_code(
        doc,
        "npm run local:services:llama\n"
        "npm run local:services:status\n"
        "npm run dev:server\n"
        "npm run dev:client\n"
        "npm run demo:smoke",
    )
    add_heading(doc, "B.2. Migration và kiểm thử", 2)
    add_code(
        doc,
        "$env:DATABASE_URL=\"postgresql://...\"\n"
        "npm run prisma:deploy --workspace server\n"
        "npm run test\n"
        "npm run test:e2e\n"
        "$env:LOAD_TOTAL=\"50\"\n"
        "$env:LOAD_CONCURRENCY=\"10\"\n"
        "npm run load:judge0",
    )

    add_heading(doc, "PHỤ LỤC C. TÀI KHOẢN DEMO", 1, page_break=True)
    add_table(
        doc,
        "Tài khoản phục vụ trình bày",
        ["Vai trò", "Email", "Mật khẩu", "Mục đích"],
        [
            ["Admin", "admin@lpp.local", "admin12345", "Quản trị học liệu, contest, reward, retention"],
            ["Nguy cơ cao", "roi.nhip@lpp.local", "hocvien123", "Biểu đồ giảm và gói cứu nhịp"],
            ["Cần theo dõi", "can.theo.doi@lpp.local", "hocvien123", "Nhóm WATCH"],
            ["Ổn định", "giu.nhip@lpp.local", "hocvien123", "Nhóm ON_TRACK"],
        ],
    )

    add_heading(doc, "KẾT LUẬN", 1, page_break=True)
    add_paragraph(
        doc,
        "CodeLearn đã hoàn thành một hệ thống học lập trình tích hợp từ học liệu đến vận hành: chấm code qua Judge0 local, AI Tutor có RAG và guardrail, thanh toán tự động, thi đua có thưởng, hồ sơ lỗi Code DNA, điểm giữ nhịp có lịch sử và dashboard can thiệp sớm. Điểm khác biệt của đề tài nằm ở vòng kín dữ liệu - giải thích - nhiệm vụ - thông báo - đo kết quả, không nằm ở tuyên bố rằng từng tính năng riêng lẻ chưa từng tồn tại.",
    )
    add_paragraph(
        doc,
        "Nhóm em đánh giá sản phẩm đã vượt mức CRUD và đủ cơ sở để triển khai thử nghiệm. Tuy nhiên, trước khi phục vụ quy mô lớn, nhóm cần đo tải trên server thật, chuyển queue/cache/limiter sang Redis dùng chung, thiết lập backup/monitoring và thu thập dữ liệu 4-8 tuần để kiểm định hiệu quả giữ chân.",
    )

    add_heading(doc, "TÀI LIỆU THAM KHẢO", 1, page_break=True)
    references = [
        "[1] React Documentation, https://react.dev/.",
        "[2] TypeScript Handbook, https://www.typescriptlang.org/docs/.",
        "[3] Express Documentation, https://expressjs.com/.",
        "[4] Prisma ORM Documentation, https://www.prisma.io/docs/.",
        "[5] PostgreSQL Documentation, https://www.postgresql.org/docs/.",
        "[6] Judge0 CE Documentation and source code, https://github.com/judge0/judge0.",
        "[7] llama.cpp server documentation, https://github.com/ggml-org/llama.cpp.",
        "[8] Qwen2.5 model documentation, https://huggingface.co/Qwen.",
        "[9] OWASP Application Security Verification Standard, https://owasp.org/www-project-application-security-verification-standard/.",
        "[10] Docker Compose Documentation, https://docs.docker.com/compose/.",
        "[11] Redis Documentation, https://redis.io/docs/.",
        "[12] Playwright Documentation, https://playwright.dev/docs/intro.",
        "[13] SePay Webhook Documentation, https://sepay.vn/.",
    ]
    for reference in references:
        add_paragraph(doc, reference, indent=False)


def build_report() -> None:
    if not TEMPLATE.exists():
        raise FileNotFoundError(f"Không tìm thấy file khung: {TEMPLATE}")
    if not REPORT.exists():
        raise FileNotFoundError(f"Không tìm thấy báo cáo cũ để lấy hình: {REPORT}")

    with tempfile.TemporaryDirectory(prefix="codelearn-report-") as temp_name:
        temp = Path(temp_name)
        assets = {path.name: path for path in ASSET_DIR.glob("*.png")}
        retention_formula, competition_formula = create_formula_figures(temp)
        screenshots = sorted((ROOT / "test-results").glob("**/retention.png"))
        desktop_screenshots = [path for path in screenshots if "desktop" in str(path).lower()]
        retention_screenshot = (
            desktop_screenshots[0]
            if desktop_screenshots
            else screenshots[0]
            if screenshots
            else assets["retention-ui.png"]
        )

        doc = Document(TEMPLATE)
        clear_body(doc)
        configure_document(doc)
        add_cover(doc)
        add_front_matter(doc)
        add_problem_section(doc, assets, retention_formula, competition_formula)
        add_business_section(doc, assets)
        add_requirement_section(doc, assets)
        add_design_section(doc, assets, retention_screenshot)
        add_appendices(doc)

        core = doc.core_properties
        core.title = "Tài liệu thiết kế chi tiết hệ thống CodeLearn"
        core.subject = "Đồ án chuyên ngành - Website hỗ trợ học lập trình"
        core.author = "Nhóm sinh viên thực hiện CodeLearn"
        core.keywords = "CodeLearn, SDD, Judge0, Llama, RAG, retention, programming education"
        core.comments = "Cập nhật theo khung tài liệu ngày 08/08/2026."

        doc.save(TEMP_REPORT)
        shutil.move(TEMP_REPORT, REPORT)
        print(f"Đã tạo: {REPORT}")
        print(f"Số đoạn: {len(doc.paragraphs)}; số bảng: {len(doc.tables)}; số hình: {len(doc.inline_shapes)}")


if __name__ == "__main__":
    build_report()
