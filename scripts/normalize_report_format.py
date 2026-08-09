from __future__ import annotations

import argparse
import re
import shutil
import zipfile
from pathlib import Path

from lxml import etree
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BLACK = "000000"
WHITE = "FFFFFF"
CAPTION_PATTERN = re.compile(r"^(?:B\u1ea3ng|H\u00ecnh)\s+\d+\s*:", re.IGNORECASE)
WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
CONTENT_TYPE_NS = "http://schemas.openxmlformats.org/package/2006/content-types"


def paragraph_style(doc: Document, name: str):
    return next(
        style
        for style in doc.styles
        if style.type == WD_STYLE_TYPE.PARAGRAPH and style.name.casefold() == name.casefold()
    )


def set_run_font(run, *, size: float | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Times New Roman")
    run.font.color.rgb = RGBColor.from_string(BLACK)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str = WHITE) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        border = borders.find(tag)
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), BLACK)


def set_table_width(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "5000")
    tbl_w.set(qn("w:type"), "pct")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "autofit")


def format_row(row, *, is_header: bool) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "true")

    if is_header:
        tbl_header = tr_pr.find(qn("w:tblHeader"))
        if tbl_header is None:
            tbl_header = OxmlElement("w:tblHeader")
            tr_pr.append(tbl_header)
        tbl_header.set(qn("w:val"), "true")

    for cell in row.cells:
        set_cell_shading(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.05
            paragraph.paragraph_format.widow_control = True
            if is_header:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, size=11, bold=True if is_header else run.bold)


def configure_styles(doc: Document) -> None:
    normal = paragraph_style(doc, "Normal")
    normal.font.name = "Times New Roman"
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.line_spacing = 1.3
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.widow_control = True

    for name, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13)):
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.italic = False
        style.font.color.rgb = RGBColor.from_string(BLACK)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True
        style.paragraph_format.widow_control = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    for style in doc.styles:
        if style.type not in (WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER):
            continue
        style.font.color.rgb = RGBColor.from_string(BLACK)
        if style.name.lower().startswith("toc") or style.name.lower() in ("hyperlink", "followedhyperlink"):
            style.font.name = "Times New Roman"
            style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Times New Roman")


def configure_pages(doc: Document) -> None:
    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.2)
        section.right_margin = Cm(2.0)
        section.header_distance = Cm(1.2)
        section.footer_distance = Cm(1.2)
        section.different_first_page_header_footer = True

        for paragraph in section.first_page_header.paragraphs + section.first_page_footer.paragraphs:
            paragraph.clear()

        for paragraph in section.header.paragraphs + section.footer.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=9 if paragraph in section.header.paragraphs else 10)


def format_paragraphs(doc: Document) -> None:
    body_started = False
    normal_style = paragraph_style(doc, "Normal")

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        style_name = paragraph.style.name

        if style_name.startswith("Heading") and not text:
            paragraph.style = normal_style
            style_name = paragraph.style.name

        if text == "L\u1edcI CAM \u0110OAN":
            body_started = True

        is_heading = style_name.startswith("Heading")
        is_caption = bool(CAPTION_PATTERN.match(text))
        has_drawing = bool(paragraph._p.xpath(".//w:drawing|.//w:pict"))

        paragraph.paragraph_format.widow_control = True
        if is_heading:
            level = int(style_name.rsplit(" ", 1)[-1])
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
            paragraph.paragraph_format.page_break_before = False
            if level == 1:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, size={1: 16, 2: 14, 3: 13}.get(level, 13), bold=True, italic=False)
        elif is_caption:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(5)
            for run in paragraph.runs:
                set_run_font(run, size=12, bold=True, italic=False)
        elif has_drawing:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
        else:
            paragraph.paragraph_format.line_spacing = 1.3
            paragraph.paragraph_format.space_after = Pt(6)
            for run in paragraph.runs:
                if run.font.name == "Consolas":
                    run.font.color.rgb = RGBColor.from_string(BLACK)
                else:
                    set_run_font(run, size=13 if body_started else None)


def format_tables(doc: Document) -> None:
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        set_table_width(table)
        set_table_borders(table)
        for index, row in enumerate(table.rows):
            format_row(row, is_header=index == 0)


def force_black_text(doc: Document) -> None:
    roots = [doc.element.body]
    for section in doc.sections:
        roots.extend(
            (
                section.header._element,
                section.footer._element,
                section.first_page_header._element,
                section.first_page_footer._element,
            )
        )

    for root in roots:
        for color in root.xpath(".//w:color"):
            color.set(qn("w:val"), BLACK)
            for attribute in ("themeColor", "themeTint", "themeShade"):
                key = qn(f"w:{attribute}")
                if key in color.attrib:
                    del color.attrib[key]


def strip_embedded_fonts(path: Path) -> None:
    temporary = path.with_name(path.stem + ".no-fonts.tmp.docx")

    with zipfile.ZipFile(path) as source, zipfile.ZipFile(temporary, "w") as destination:
        for item in source.infolist():
            name = item.filename
            if name.startswith("word/fonts/"):
                continue

            data = source.read(name)
            if name == "word/fontTable.xml":
                root = etree.fromstring(data)
                for font in root:
                    for child in list(font):
                        if child.tag.rsplit("}", 1)[-1].startswith("embed"):
                            font.remove(child)
                data = etree.tostring(root, encoding="utf-8", xml_declaration=True, standalone=True)
            elif name == "word/_rels/fontTable.xml.rels":
                root = etree.fromstring(data)
                for relationship in list(root):
                    if relationship.get("Type", "").endswith("/font"):
                        root.remove(relationship)
                data = etree.tostring(root, encoding="utf-8", xml_declaration=True, standalone=True)
            elif name == "word/settings.xml":
                root = etree.fromstring(data)
                embedded_font_tags = {
                    f"{{{WORD_NS}}}embedTrueTypeFonts",
                    f"{{{WORD_NS}}}embedSystemFonts",
                    f"{{{WORD_NS}}}saveSubsetFonts",
                }
                for child in list(root):
                    if child.tag in embedded_font_tags:
                        root.remove(child)
                data = etree.tostring(root, encoding="utf-8", xml_declaration=True, standalone=True)
            elif name == "[Content_Types].xml":
                root = etree.fromstring(data)
                override_tag = f"{{{CONTENT_TYPE_NS}}}Override"
                for override in list(root.findall(override_tag)):
                    if override.get("PartName", "").startswith("/word/fonts/"):
                        root.remove(override)
                data = etree.tostring(root, encoding="utf-8", xml_declaration=True, standalone=True)

            destination.writestr(item, data)

    shutil.move(temporary, path)


def normalize(source: Path, destination: Path) -> None:
    doc = Document(source)
    configure_styles(doc)
    configure_pages(doc)
    format_paragraphs(doc)
    format_tables(doc)
    force_black_text(doc)

    destination.parent.mkdir(parents=True, exist_ok=True)
    temporary = destination.with_name(destination.stem + ".formatting.tmp.docx")
    doc.save(temporary)
    shutil.move(temporary, destination)
    strip_embedded_fonts(destination)


def main() -> None:
    parser = argparse.ArgumentParser(description="Normalize the CodeLearn report to an academic black-and-white format.")
    parser.add_argument("source", type=Path)
    parser.add_argument("destination", type=Path, nargs="?")
    args = parser.parse_args()
    normalize(args.source.resolve(), (args.destination or args.source).resolve())


if __name__ == "__main__":
    main()
